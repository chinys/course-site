# -*- coding: utf-8 -*-
"""
10종 계약서(+매뉴얼 지침)를 Word(.docx) 파일로 생성하여 static/downloads/ 에 저장합니다.
HTML 테이블을 Word 테이블로 변환합니다.
"""
import os, re, html as html_mod
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

KOREAN_FONT = '나눔고딕'
BASE_DIR = os.path.join(os.path.dirname(__file__), '..')
OUT_DIR = os.path.join(BASE_DIR, 'static', 'downloads')
os.makedirs(OUT_DIR, exist_ok=True)


def _set_kr(run, size=Pt(10), bold=False):
    """Latin + East Asian 글꼴을 모두 한글 글꼴로 설정"""
    run.font.size = size
    run.bold = bold
    run.font.name = KOREAN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), KOREAN_FONT)
    rFonts.set(qn('w:ascii'), KOREAN_FONT)
    rFonts.set(qn('w:hAnsi'), KOREAN_FONT)


def _strip(text):
    """HTML 태그 제거 + 엔티티 디코딩"""
    # HTML 내 연속된 공백이나 줄바꿈을 단일 공백으로 치환 (HTML 렌더링 방식과 동일하게)
    text = re.sub(r'\s+', ' ', text)
    # Replace <br> with newline
    text = re.sub(r'<br\s*/?>', '\n', text)
    # Replace closing tags of block elements with double newline to preserve spacing
    text = re.sub(r'</(p|div|ul|ol|h[1-6])>', '\n\n', text)
    # Replace closing tags of list items with a single newline
    text = re.sub(r'</li>', '\n', text)
    # Remove all remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Decode entities
    text = html_mod.unescape(text)
    
    # 각 줄의 앞뒤 공백 제거
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # 연속된 빈 줄을 최대 2개로 제한
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    return text


def _parse_table_rows(table_html):
    """HTML 테이블에서 행/셀 추출 → list of list of (text, is_header)"""
    rows = []
    for tr_match in re.finditer(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL):
        tr = tr_match.group(1)
        cells = []
        for cell_match in re.finditer(r'<(th|td)[^>]*>(.*?)</\1>', tr, re.DOTALL):
            tag = cell_match.group(1)
            content = _strip(cell_match.group(2))
            cells.append((content, tag == 'th'))
        if cells:
            rows.append(cells)
    return rows


def _add_word_table(doc, rows):
    """파싱된 테이블 데이터를 Word 테이블로 추가"""
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, (text, is_header) in enumerate(row_data):
            if j >= max_cols:
                break
            cell = row.cells[j]
            cell.text = ''
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            if is_header:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            _set_kr(run, size=Pt(9), bold=is_header)
    doc.add_paragraph()  # 테이블 후 빈 줄


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_kr(run, size=Pt(16 if level == 1 else 13), bold=True)


def _add_body(doc, text):
    """일반 텍스트를 문단으로 추가"""
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = Pt(16)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(16)
        if line.startswith('제') and '조' in line[:8]:
            run = p.add_run(line)
            _set_kr(run, size=Pt(11), bold=True)
        elif line.startswith('부 칙') or line.startswith('부칙'):
            run = p.add_run(line)
            _set_kr(run, size=Pt(11), bold=True)
        else:
            run = p.add_run(line)
            _set_kr(run, size=Pt(10))


def _add_html_content(doc, html_content):
    """HTML 콘텐츠를 테이블과 텍스트로 분리하여 Word에 추가"""
    # 테이블과 비테이블 구간 분리
    parts = re.split(r'(<table[^>]*>.*?</table>)', html_content, flags=re.DOTALL)
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if part.startswith('<table'):
            rows = _parse_table_rows(part)
            _add_word_table(doc, rows)
        else:
            text = _strip(part)
            if text:
                _add_body(doc, text)


def _init_doc():
    """한글 글꼴이 설정된 Document 생성"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = KOREAN_FONT
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), KOREAN_FONT)
    rFonts.set(qn('w:ascii'), KOREAN_FONT)
    rFonts.set(qn('w:hAnsi'), KOREAN_FONT)
    return doc


def _build_docx(title, contract_html, manual_html, filename):
    """하나의 계약서 + 매뉴얼 지침을 DOCX로 저장"""
    doc = _init_doc()
    _add_heading(doc, title, level=1)
    _add_html_content(doc, contract_html)

    if manual_html:
        doc.add_page_break()
        _add_heading(doc, '매뉴얼 지침 및 원가 분류 유의사항', level=2)
        _add_html_content(doc, manual_html)

    doc.save(os.path.join(OUT_DIR, filename))
    print(f'  [OK] {filename} 생성 완료')


def _split_contract_and_manual(block_html):
    """계약서 div (bg-gray-50)과 매뉴얼 div (💡) 분리"""
    # 1. 계약서 부분 찾기
    contract = ''
    marker = 'bg-gray-50 border border-gray-300'
    idx = block_html.find(marker)
    if idx >= 0:
        tag_end = block_html.find('>', idx)
        if tag_end >= 0:
            rest = block_html[tag_end + 1:]
            
            # 매뉴얼 지침 시작 부분 찾기 (💡 아이콘이 있는 div 기준)
            manual_start_marker = '💡'
            manual_idx = rest.find(manual_start_marker)
            
            if manual_idx >= 0:
                # 💡 이전까지가 계약서 본문 (보통 </div> </div> 가 사이에 있음)
                contract = rest[:manual_idx]
            else:
                contract = rest

    # 2. 매뉴얼 부분 찾기
    manual = ''
    manual_idx = block_html.find('💡')
    if manual_idx >= 0:
        # 💡 이후 텍스트 중 실제 '매뉴얼 지침' 텍스트 시작점을 찾음
        title_idx = block_html.find('매뉴얼 지침', manual_idx)
        if title_idx >= 0:
            manual = block_html[title_idx:]
            
    return contract, manual


def generate_all():
    import sys
    sys.path.insert(0, os.path.join(BASE_DIR, 'scripts'))
    from _section_a import section_a
    from _section_b import section_b
    from _section_c import section_c
    from _section_d import section_d
    from _section_e import section_e

    sa, sb, sc, sd, se = section_a(), section_b(), section_c(), section_d(), section_e()

    contracts = [
        ('A1', '물품공급 계약서(자재포함 외주) (A-1)', sa, '<!-- A-1', '<!-- A-2'),
        ('A2', '외주 사급가공 계약서(자재 사급 외주) (A-2)', sa, '<!-- A-2', '<!-- A-3'),
        ('A3', '사내가공제작(사내도장) 계약서 (A-3)', sa, '<!-- A-3', '<!-- A-4'),
        ('A4', '물품공급(건설현장 납품용) 계약서 (A-4)', sa, '<!-- A-4', None),
        ('B1', '건설기계 장비 임대차 계약서 (B-1)', sb, '<!-- B-1', '<!-- B-2'),
        ('B2', '가설재(비계·서포트) 임대차계약서 (B-2)', sb, '<!-- B-2', None),
        ('C1', '기술·품질용역 위탁계약서(NDT) (C-1)', sc, '<!-- C-1', '<!-- C-2'),
        ('C2', '건설현장 급식(함바식당) 운영 위탁계약서 (C-2)', sc, '<!-- C-2', None),
        ('D1', '전문 기술자문 및 현장관리 위촉계약서 (D-1)', sd, '<!-- D-1', '<!-- D-2'),
        ('D2', '일용직 근로계약서 (D-2)', sd, '<!-- D-2', None),
        ('E1', '화물(강재) 운반 용역계약서 (E-1)', se, '<!-- E-1', None),
    ]

    print('DOCX 계약서 생성 시작...\n')
    for cid, title, src, comment, next_comment in contracts:
        start_idx = src.find(comment)
        if start_idx < 0:
            print(f'  [실패] {cid} 블록을 찾을 수 없습니다.')
            continue
        if next_comment:
            end_idx = src.find(next_comment, start_idx + len(comment))
            block = src[start_idx:end_idx] if end_idx > 0 else src[start_idx:]
        else:
            block = src[start_idx:]
        contract_part, manual_part = _split_contract_and_manual(block)
        _build_docx(title, contract_part, manual_part, f'contract_{cid}.docx')

    print(f'\n[성공] 모든 계약서 DOCX 생성 완료! -> {OUT_DIR}')


if __name__ == '__main__':
    generate_all()
